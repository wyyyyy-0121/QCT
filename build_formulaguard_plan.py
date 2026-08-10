# -*- coding: utf-8 -*-
"""Build the FormulaGuard competition research plan as a polished DOCX."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "deliverables"
OUT_PATH = OUT_DIR / "FormulaGuard_银奖目标研究计划书.docx"

# Selected design preset: narrative_proposal.
# Named override: Microsoft YaHei replaces Calibri to ensure stable CJK rendering.
FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Aptos"
FONT_MATH = "Cambria Math"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5F6B7A"
LIGHT = "F4F6F9"
BLUE_GRAY = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
GOLD = "7A5A00"
PALE_GOLD = "FFF8E8"
GREEN = "276749"
PALE_GREEN = "EDF7F1"
RED = "9B1C1C"
PALE_RED = "FFF1F1"
BORDER = "CBD5E1"
TABLE_WIDTH = 9360
TABLE_INDENT = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name=FONT_CN, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGINS):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.find(qn("w:tcMar"))
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for key in ("top", "start", "bottom", "end"):
        node = tcmar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tcmar.append(node)
        node.set(qn("w:w"), str(margins[key]))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6, inside=True):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    edges = ["top", "left", "bottom", "right"]
    if inside:
        edges += ["insideH", "insideV"]
    for edge in edges:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths: Sequence[int], indent=TABLE_INDENT):
    if sum(widths) != TABLE_WIDTH:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH}, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(TABLE_WIDTH))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trpr.append(cant)
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def set_cell_text(cell, text, *, bold=False, color=NAVY, size=9.3,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int],
              header_fill=LIGHT, font_size=9.1, zebra=False):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], header_fill)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=NAVY, size=9.2,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for r_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        if zebra and r_idx % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "FAFBFC")
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value, size=font_size)
    # Reapply geometry after rows have been added so every tcW is explicit.
    set_table_geometry(table, widths)
    return table


def add_callout(doc, title, body, *, fill=LIGHT, title_color=BLUE, body_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH])
    set_table_borders(table, color=fill, size=4, inside=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=title_color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(1)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color=body_color)
    return table


def add_spacer(doc, pts):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    return p


def add_body(doc, text, *, bold_prefix=None, color=NAVY, size=11, after=8,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, keep=False):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=size, color=color, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=size, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color)
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run_font(r, name=FONT_MATH, size=10.8, color=NAVY)
    return p


def _next_abstract_num_id(numbering):
    ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    return max(ids, default=-1) + 1


def _next_num_id(numbering):
    ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    return max(ids, default=0) + 1


def make_numbering(doc, kind="bullet"):
    numbering = doc.part.numbering_part.element
    abstract_id = _next_abstract_num_id(numbering)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    ml = OxmlElement("w:multiLevelType")
    ml.set(qn("w:val"), "singleLevel")
    abstract.append(ml)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(fmt)
    txt = OxmlElement("w:lvlText")
    txt.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(txt)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "279")
    ppr.append(ind)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT_CN)
    rpr.append(rfonts)
    lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num_id = _next_num_id(numbering)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def add_list(doc, items: Iterable[str], *, numbered=False, color=NAVY, size=10.7):
    num_id = make_numbering(doc, "decimal" if numbered else "bullet")
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        ppr = p._p.get_or_add_pPr()
        numpr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numid = OxmlElement("w:numId")
        numid.set(qn("w:val"), str(num_id))
        numpr.append(ilvl)
        numpr.append(numid)
        ppr.append(numpr)
        r = p.add_run(item)
        set_run_font(r, size=size, color=color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = False
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}[level],
                 color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = paragraph.add_run("第 ")
    set_run_font(r1, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = OxmlElement("w:r")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_text)
    r.append(fld_end)
    paragraph._p.append(r)
    r2 = paragraph.add_run(" 页")
    set_run_font(r2, size=9, color=MUTED)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        st = styles[f"Heading {level}"]
        st.font.name = FONT_CN
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.widow_control = True

    caption = styles["Caption"]
    caption.font.name = FONT_CN
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)


def setup_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def setup_header_footer(section):
    # Running header: quiet two-column label, no bottom border.
    header = section.header
    p0 = header.paragraphs[0]
    p0._element.getparent().remove(p0._element)
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [4680, 4680], indent=0)
    # Remove any borders.
    tblpr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblpr.append(borders)
    set_cell_text(table.cell(0, 0), "FormulaGuard 研究计划书", size=8.5, color=MUTED)
    set_cell_text(table.cell(0, 1), "2026 丘成桐中学科学奖 · 计算机", size=8.5, color=MUTED,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)

    fp_header = section.first_page_header
    fp_header.paragraphs[0].text = ""

    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_page_number(p)
    section.first_page_footer.paragraphs[0].text = ""


def add_cover(doc):
    add_spacer(doc, 38)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("2026 丘成桐中学科学奖（计算机奖）")
    set_run_font(r, size=11, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run("FormulaGuard 研究计划书")
    set_run_font(r, size=29, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("无输出真值条件下基于图干预的电子表格静默公式根因定位")
    set_run_font(r, size=15, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("Oracle-Free Root-Cause Localization of Silent Spreadsheet Formula Errors via Graph Interventions")
    set_run_font(r, name=FONT_LATIN, size=10.5, color=MUTED, italic=True)

    table = doc.add_table(rows=4, cols=4)
    set_table_geometry(table, [1280, 3400, 1280, 3400])
    set_table_borders(table, color=LIGHT, size=4)
    meta = [
        ("参赛学生", "________________", "学校", "________________"),
        ("指导教师", "________________", "研究方向", "电子表格程序分析"),
        ("计划周期", "2026.08.10—09.15", "目标层级", "本科研究入门水平"),
        ("主目标", "冲击银奖", "文档版本", "v1.0 · 2026.08.10"),
    ]
    for i, row in enumerate(meta):
        for j, value in enumerate(row):
            label = j in (0, 2)
            if label:
                set_cell_shading(table.rows[i].cells[j], LIGHT)
            set_cell_text(table.rows[i].cells[j], value, bold=label, color=BLUE if label else NAVY,
                          size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER if label else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, [1280, 3400, 1280, 3400])
    add_spacer(doc, 20)
    add_callout(
        doc,
        "计划定位",
        "这是一个“能真正跑起来”的实验型研究计划：最终成果必须包含可复现代码、自动生成的带标签错误基准、完整对比实验、真实表格验证和可演示原型。计划中的数值均为预先设定的目标，不是已经得到的实验结果。",
        fill=PALE_GOLD,
        title_color=GOLD,
    )
    add_spacer(doc, 26)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("银奖导向 · 证据优先 · 范围可控 · 结果诚实")
    set_run_font(r, size=10.5, color=BLUE, bold=True)
    doc.add_page_break()


def add_flow_table(doc, labels):
    widths = [1500, 465, 1500, 465, 1500, 465, 1500, 465, 1500]
    table = doc.add_table(rows=1, cols=9)
    set_table_geometry(table, widths)
    # No outer grid; boxes are individually shaded.
    tblpr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblpr.append(borders)
    for i in range(9):
        if i % 2 == 0:
            idx = i // 2
            set_cell_shading(table.cell(0, i), BLUE_GRAY if idx % 2 == 0 else LIGHT)
            set_cell_text(table.cell(0, i), labels[idx], bold=True, size=9.1,
                          color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            set_cell_text(table.cell(0, i), "→", bold=True, size=13, color=BLUE,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, widths)
    return table


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    setup_page(section)
    setup_header_footer(section)

    props = doc.core_properties
    props.title = "FormulaGuard 银奖目标研究计划书"
    props.subject = "无输出真值条件下基于图干预的电子表格静默公式根因定位"
    props.author = "参赛学生（待填写）"
    props.keywords = "FormulaGuard, spreadsheet, formula error, fault localization, graph intervention"
    props.comments = "narrative_proposal preset; CJK font override: Microsoft YaHei"

    add_cover(doc)

    add_heading(doc, "阅读导航", 1)
    add_body(doc, "本计划按“为什么值得做—做什么—怎样编码—怎样实验—如何按期交付”的顺序组织。真正决定项目质量的不是题目听起来多先进，而是错误标签是否可靠、对比是否公平、实验是否可复现、结论是否经得起追问。")
    nav_rows = [
        ("A. 项目定义", "第1—4节", "定位、问题、假设、边界与创新"),
        ("B. 技术方案", "第5—7节", "系统流程、GIR算法、候选修复"),
        ("C. 实验核心", "第8—12节", "基准、仿真、对比、消融、统计"),
        ("D. 执行管理", "第13—17节", "代码结构、进度、风险、银奖门槛"),
        ("E. 写作与合规", "第18—20节", "论文、展示、AI披露与参考资料"),
        ("F. 附录模板", "附录A—C", "错误记录、运行清单、研究日志"),
    ]
    add_table(doc, ["模块", "对应部分", "阅读重点"], nav_rows, [1900, 1800, 5660], zebra=True)
    add_spacer(doc, 7)
    add_callout(
        doc,
        "先回答你的问题：要不要做实验？",
        "必须做，而且实验是整份作品的核心证据。至少要在本地运行解析器、依赖图构建、错误注入、公式重算、定位排序、候选修复和统计分析；如果只有方案文字而没有代码、数据和结果，无法支撑银奖目标。",
        fill=PALE_GREEN,
        title_color=GREEN,
    )
    doc.add_page_break()

    add_heading(doc, "执行摘要", 1)
    add_body(doc, "FormulaGuard 研究大型 Excel 表格中的“静默公式错误”：公式能够正常计算，不显示 #REF! 或 #VALUE!，但因为引用、范围、运算符或复制偏移写错，结果在依赖链中持续传播。普通用户往往先看到下游汇总异常，却难以找到最初写错的源单元格。本项目希望在不给算法正确输出、测试用例或人工标注的条件下，自动给出最可疑的源错误单元格及候选修复。")
    add_body(doc, "计划的核心不是发明依赖图本身，而是提出一个清楚、可检验的问题设定：无输出真值（oracle-free）、静默错误（silent error）、根因定位（root-cause localization）。方法上，先把公式归一化并分组，再构建单元格依赖图；随后对可疑单元格进行候选公式替换，观察替换是否同时降低公式模式、局部图结构和下游行为的不一致，形成“图干预责任分数”（Graph Intervention Responsibility, GIR）。")
    add_body(doc, "实验上，项目将自动生成传播深度可控的 PropagationBench 基准，并在真实表格上做外部验证。所有方法只看到带错工作簿；无错版本和源错误标签只用于评分。主实验比较 FormulaGuard 与随机排序、图异常、公式模式多数法、Excel 不一致公式提示，以及可复现的 ExceLint-like/WARDER-like 方法；SFL-Oracle 作为使用输出真值的上界，不混入同一公平组。")

    summary_rows = [
        ("核心问题", "不提供正确输出时，能否定位造成多个下游异常的最初公式错误？"),
        ("主要原创点", "问题定义 + GIR 图干预分数 + 传播深度可控基准"),
        ("实现强度", "Python 可复现原型；自动错误注入；批量实验；Top-5解释报告"),
        ("主评价", "Top-1/3/5、MRR、EXAM、修复命中率、运行时间、干净表误报"),
        ("工作量边界", "单源错误为主；常见函数；无宏/外链/循环引用/动态数组"),
        ("成功判据", "相对最强无真值基线有稳定提升，并在真实表格上保持可解释性"),
    ]
    add_table(doc, ["项目要素", "约定"], summary_rows, [1900, 7460], zebra=True)
    add_spacer(doc, 6)
    add_callout(doc, "获奖预期说明", "这条路线具备参赛和冲击银奖的合理性，但没有任何方案能保证奖项。评审最终看的是原创边界、实验质量、真实提升、答辩可信度和学生本人对代码与论文的掌握程度。", fill=PALE_RED, title_color=RED)

    add_heading(doc, "1. 项目定位与目标", 1)
    add_heading(doc, "1.1 目标层级", 2)
    add_body(doc, "项目按“优秀本科课程研究/本科科研入门”的强度设计：问题要正式、算法要可实现、实验要完整，但不追求覆盖所有 Excel 语法，也不要求提出博士论文级别的统一理论。高中生应能够解释每个模块、运行每个实验、复核每张结果表，并在答辩时现场演示一个真实案例。")
    add_heading(doc, "1.2 银奖导向的四条主线", 2)
    add_list(doc, [
        "清楚：把研究对象限定为无真值条件下的单源静默公式根因定位，避免题目无限扩张。",
        "原创：将“候选修复带来的全局不一致下降”正式定义为图干预责任，而不是简单拼接现有异常分数。",
        "可信：用自动生成且可验证的错误基准、模板隔离的数据划分和真实表格外部验证形成证据链。",
        "可展示：最终工具一键输出 Top-5 可疑单元格、候选修复、影响路径和诊断报告。",
    ])
    add_heading(doc, "1.3 不把项目做成什么", 2)
    add_table(doc, ["不做", "原因", "替代方案"], [
        ("覆盖全部 Excel 函数和插件", "五周内不可验证，容易产生大量边界错误", "先支持常用算术、SUM/AVERAGE/MIN/MAX 与直接引用"),
        ("把大语言模型作为主算法", "难以稳定、难以解释，也容易削弱学生原创性", "核心算法完全可复现；LLM仅可用于辅助说明或候选扩展"),
        ("一次处理多个源错误", "评价和因果归因难度显著上升", "主实验单源错误；多错误仅作探索性附录"),
        ("宣称依赖图或公式聚类是首创", "这些思想已有充分先行研究", "把原创性放在问题设定、GIR和传播基准上"),
        ("为了达到目标筛掉失败样本", "会造成选择偏差，答辩时难以自证", "预先登记规则；所有有效样本按固定脚本统计"),
    ], [1700, 3350, 4310], zebra=True)

    add_heading(doc, "2. 研究背景与意义", 1)
    add_body(doc, "电子表格把程序逻辑分散在成百上千个单元格中。一个源公式写错后，依赖它的汇总、图表和决策指标可能全部改变。最危险的并不是显式错误，而是“看起来正常”的数值：工作簿能打开、公式能计算、结果有合理数量级，却已偏离真实业务逻辑。")
    add_body(doc, "传统调试通常依赖用户指出正确结果或失败测试；但真实办公环境常常没有这种输出真值。另一类方法通过公式模式或空间异常寻找离群公式，适合发现局部不一致，但不一定能区分“最初写错的源头”和“被错误传播影响的下游症状”。因此，本研究关心的是：在没有正确答案的前提下，能否利用公式族、依赖结构和反事实替换后的系统变化，对源错误承担的责任进行排序。")
    add_callout(doc, "实际价值", "如果工具能把需要人工检查的范围从数百个公式压缩到 Top-5，并说明“为何可疑、影响了谁、怎样修复”，它就能显著降低审计和排错成本。", fill=LIGHT, title_color=BLUE)

    add_heading(doc, "3. 相关工作与原创边界", 1)
    add_body(doc, "电子表格调试并不是空白领域。已有研究覆盖目标驱动修复、数据流分析、基于测试的故障定位、公式气味、公式聚类、空间异常检测、神经符号修复和蜕变测试等方向。[1]—[12] 因此，参赛论文必须主动说明“借鉴了什么、不同在哪里、为什么仍值得研究”。")
    prior_rows = [
        ("GoalDebug (2007)", "依据期望输出和用户目标生成公式修复", "需要目标/预期结果；本项目不提供输出真值"),
        ("Spreadsheet Debugging (2008)", "数据流与影响传播分析", "依赖图是基础；本项目增加候选替换后的责任度量"),
        ("Hofer等 (2013)", "SFL、切片、约束调试的实证比较", "SFL依赖测试通过/失败信息；本项目将其作为有真值上界"),
        ("FaultySheet Detective (2014)", "公式气味与SFL联合定位", "本项目主比较组不使用测试结果"),
        ("MUSSCO (2015)", "基于变异的公式修复候选", "借鉴小编辑候选；评分改为无真值图干预"),
        ("ExceLint (2018)", "信息论与空间公式异常", "强无真值对照；本项目强调根因和传播"),
        ("Metric-based (2019)", "使用表格度量预测故障倾向", "更偏静态风险预测；本项目针对具体源错误排名"),
        ("WARDER (2020)", "公式族/聚类辅助错误发现与修复", "借鉴公式族；加入依赖图责任与深度实验"),
        ("LaMirage (2022)", "神经符号低代码公式修复", "模型规模和任务不同；本项目优先可解释与CPU可运行"),
        ("FLAME/RING (2023)", "预训练模型驱动公式生成/修复", "作为背景，不作为必须复现的主基线"),
        ("FoRepBench (2025)", "上下文公式修复评测样本", "偏修复；可用于候选生成的补充验证"),
        ("SEDMR (2026)", "用蜕变关系检测电子表格缺陷", "借鉴无真值行为一致性思想；主任务改为源定位"),
    ]
    add_table(doc, ["代表工作", "已有贡献", "FormulaGuard 的区别"], prior_rows, [2100, 3100, 4160], font_size=8.45, zebra=True)
    add_body(doc, "表1 相关工作用于界定原创边界；完整出处见第20节。", size=9, color=MUTED, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "3.1 本项目可以主张的三项贡献", 2)
    add_list(doc, [
        "问题贡献：形式化“无输出真值、静默错误、单源根因定位”的任务、输入、输出和评价协议。",
        "方法贡献：提出图干预责任分数 GIR，把局部公式异常、图结构异常、下游一致性改善和编辑成本统一到候选干预框架。",
        "实验贡献：建立传播深度可控、标签可追溯的 PropagationBench，并展示性能随传播深度和错误类型的变化。",
    ])
    add_heading(doc, "3.2 论文中必须避免的表述", 2)
    add_list(doc, [
        "避免“首次使用依赖图定位 Excel 错误”；正确说法是“在既有依赖分析基础上提出新的无真值干预责任评分”。",
        "避免“全面优于现有方法”；正确说法是“在本研究定义的数据范围和公平设置中，相对指定基线获得提升”。",
        "避免“自动修复所有公式错误”；正确说法是“给出 Top-k 根因与有限候选修复，最终由用户确认”。",
    ])

    add_heading(doc, "4. 研究问题、假设与范围", 1)
    add_heading(doc, "4.1 研究问题（Research Questions）", 2)
    rq_rows = [
        ("RQ1", "无输出真值时，GIR 能否比纯公式模式或纯图分数更准确地定位源错误？", "Top-k、MRR、EXAM"),
        ("RQ2", "传播深度增加时，各方法的性能如何变化？", "按浅/中/深分层结果"),
        ("RQ3", "哪些组件真正贡献提升？", "消融实验与效应量"),
        ("RQ4", "候选修复是否能覆盖并排到真实公式？", "Coverage、Repair Top-k"),
        ("RQ5", "方法在真实表格、干净表格和较大表格上是否仍可靠？", "外部验证、误报率、运行时间"),
    ]
    add_table(doc, ["编号", "问题", "主要证据"], rq_rows, [900, 5960, 2500], zebra=True)
    add_heading(doc, "4.2 可证伪假设", 2)
    add_list(doc, [
        "H1：在相同无真值输入下，完整 GIR 的 Top-1 和 MRR 显著高于最强无真值基线。",
        "H2：只用节点下游规模会偏爱上游大节点；加入候选干预后的不一致下降能减少这种误排。",
        "H3：传播越深，症状越多，单纯空间异常方法越容易失效；GIR 的相对优势会更明显。",
        "H4：公式族多数模板可使真实修复进入候选集的比例达到可用水平，但对唯一公式区域效果较弱。",
    ])
    add_heading(doc, "4.3 输入、输出和约束", 2)
    scope_rows = [
        ("输入", "一个 .xlsx 工作簿；算法不知道正确版本、正确输出和源错误位置"),
        ("输出", "Top-5 可疑公式单元格、GIR分数、候选修复、影响路径和理由"),
        ("主错误模型", "一次只注入一个源公式错误；错误结果可正常计算，无显式Excel错误码"),
        ("支持公式", "算术、比较、直接引用、SUM、AVERAGE、MIN、MAX；逐步扩展 IF/COUNTIF"),
        ("排除", "VBA宏、外部链接、循环引用、动态数组、易变函数、多源并发错误"),
        ("适用表格", "存在重复行/列/块结构，能够形成可比较公式族的工作簿"),
        ("真值隔离", "正确版本只在基准生成和最终评分脚本中可见，定位模块不得读取"),
    ]
    add_table(doc, ["项目", "约定"], scope_rows, [1900, 7460], zebra=True)

    doc.add_page_break()
    add_heading(doc, "5. 总体技术路线", 1)
    add_body(doc, "系统由五个串联模块组成。每个模块都有独立输入输出和单元测试；这样即使最终算法效果不够理想，也能定位问题出在公式解析、候选覆盖、依赖图还是评分设计，而不是只得到一个无法解释的总分。")
    add_flow_table(doc, ["公式解析\n与归一化", "公式族\n与局部模式", "依赖图\n与传播路径", "候选替换\n反事实干预", "GIR排序\n与诊断报告"])
    add_body(doc, "图1 FormulaGuard 主流程。算法输入为带错工作簿；正确工作簿仅存在于独立评测端。", size=9, color=MUTED, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "5.1 公式解析与归一化", 2)
    add_body(doc, "将每个公式解析为抽象语法树（AST），把绝对地址和相对偏移分开表示。例如，同一列中 =B2*C2 与 =B3*C3 在 R1C1 相对形式下属于同一模板。归一化输出至少包含操作符序列、函数名、引用相对偏移、区域形状、绝对/相对标记和常量类型。")
    add_list(doc, [
        "单元测试覆盖括号优先级、跨表引用、范围、绝对引用、混合引用和错误输入。",
        "解析失败时记录原因并跳过，而不是悄悄生成错误依赖边。",
        "同一公式的A1与R1C1表达应可往返验证，避免复制偏移判断错误。",
    ])
    add_heading(doc, "5.2 公式族与局部图模式", 2)
    add_body(doc, "在相邻行、相邻列或重复块中，根据归一化公式指纹聚类形成公式族。对每个公式单元格，再提取局部图模式：入度、出度、引用列类型、范围形状、同族邻居、下游深度和可达汇总节点。公式异常与图异常分别计算，避免一个指标同时承担所有解释。")

    add_heading(doc, "6. 图干预责任分数（GIR）", 1)
    add_heading(doc, "6.1 不一致能量", 2)
    add_body(doc, "设工作簿为 W，定义三个无需正确输出即可计算的分量：公式族不一致 E_formula、局部依赖图模式不一致 E_graph、以及在允许的输入扰动或候选替换下的下游行为不一致 E_behavior。")
    add_equation(doc, "E(W) = w_f · E_formula(W) + w_g · E_graph(W) + w_b · E_behavior(W)")
    add_body(doc, "权重 w_f、w_g、w_b 不在测试集上反复调节。主方案使用开发集网格搜索或固定为等权；最终论文同时报告等权结果和开发集选择结果，以防调参造成夸大。")

    add_heading(doc, "6.2 候选干预", 2)
    add_body(doc, "对于候选单元格 v，从同族邻居平移、行/列多数模板、相邻块对应位置和一次AST小编辑中生成不超过 K 个候选公式 C(v)。把 v 暂时替换为候选 f 后，仅重算受影响的下游子图，得到能量下降：")
    add_equation(doc, "ΔE(v, f) = E(W) − E(W[v ← f])")
    add_body(doc, "如果替换后多个层面的不一致同时下降，说明 v 更可能是源错误；如果只是一个上游大节点但替换不能恢复一致，则不应仅因影响范围大而排到最前。")

    add_heading(doc, "6.3 正式责任分数", 2)
    add_equation(doc, "GIR(v) = I(v) · max_{f∈C(v)} max(0, ΔE(v,f)) / (1 + EditCost(v,f))")
    add_body(doc, "其中 I(v) 是有界图影响因子，可由标准化后代数和可达汇总节点数构成；EditCost 惩罚大幅改写，优先选择平移、边界修正或单个运算符替换等小改动。输出时同时保存最优候选 f*、三个能量分量变化和影响路径，形成可解释证据。")

    add_heading(doc, "6.4 简化命题与复杂度", 2)
    add_callout(doc, "可验证命题（本科水平）", "若工作簿只含一个源错误；该单元格所在公式族的多数模板正确；候选集中包含正确公式；且正确替换是唯一使总不一致能量最小的替换，则源错误的干预能量下降严格大于同族非错误单元格。论文可给出一页证明，并用反例说明每个条件为何必要。", fill=LIGHT, title_color=DARK_BLUE)
    add_body(doc, "若公式数为 F、引用数为 R、图含 V 个节点和 E 条边，解析和建图约为 O(F+R)。对 M 个候选单元格、每个 K 个公式做完整重算，最坏约 O(MK(V+E))；实现时缓存局部能量并只遍历下游子图，以实测运行时间说明优化效果，而不追求复杂理论。")

    add_heading(doc, "7. 候选修复与解释输出", 1)
    candidate_rows = [
        ("同列邻居平移", "把上一行/下一行同族公式翻译到当前坐标", "最常见复制偏移，编辑小"),
        ("行/列多数模板", "选择局部窗口中占多数的归一化公式", "抑制单个离群公式"),
        ("重复块对应位置", "从相邻月份/部门/实验组的同位置复制模板", "适合报表块结构"),
        ("范围边界修正", "对起止行列进行 ±1 或吸附到相邻模板", "覆盖漏一行、多一行"),
        ("AST一次编辑", "运算符、函数名或绝对引用标记的一次替换", "覆盖常见人工输入错误"),
    ]
    add_table(doc, ["来源", "生成方式", "意义"], candidate_rows, [2200, 4100, 3060], zebra=True)
    add_body(doc, "为了保证候选搜索可控，K 默认不超过 10；论文报告 K=3/5/10/20 的覆盖率—运行时间曲线。诊断报告必须区分“检测到可疑源”与“候选修复已被验证”，不得自动覆盖原文件。")
    add_callout(doc, "演示输出示例", "1) Budget!H27，GIR=0.84；2) 建议 =SUM(H18:H26)；3) 原式遗漏H26；4) 该单元格影响 14 个下游公式和 2 个汇总输出；5) 替换后公式族、依赖范围和下游一致性同时改善。", fill=PALE_GREEN, title_color=GREEN)

    doc.add_page_break()
    add_heading(doc, "8. PropagationBench：传播深度可控基准", 1)
    add_body(doc, "基准的目的不是制造大量随意错误，而是产生标签可靠、结构多样、传播深度可控、能够复算的样本。每个错误实例包含正确工作簿、带错工作簿、源单元格、正确公式、错误公式、错误类型、传播深度、受影响节点、输出差异和随机种子。定位算法只能读取带错工作簿。")

    add_heading(doc, "8.1 基础工作簿", 2)
    template_rows = [
        ("预算/费用", "逐项成本→分类小计→总计", "范围与引用偏移"),
        ("销售/月报", "明细→区域汇总→季度汇总", "深层传播与重复块"),
        ("库存", "入库/出库→结余→预警", "跨期引用与条件公式"),
        ("成绩表", "分项→总分→排名/统计", "权重、平均值和复制"),
        ("实验记录", "多次测量→均值→误差→结论", "科研情境、公式链清楚"),
        ("能耗/用水", "日数据→周/月汇总→同比", "范围边界与多层汇总"),
        ("项目排期", "工时→阶段→资源汇总", "跨列依赖与绝对引用"),
        ("发票/税费", "数量×单价→税额→应付", "运算符、绝对参数"),
    ]
    add_table(doc, ["模板族", "典型结构", "适合错误"], template_rows, [1900, 4100, 3360], zebra=True)
    add_body(doc, "计划编写 8—10 个 Python 模板生成器，每族生成 5—8 个参数化工作簿，形成 50—80 个基础工作簿。自动化可以控制规模并保留结构标签；另外保留一部分手工制作工作簿，检验算法是否只适应生成器规律。")

    add_heading(doc, "8.2 六类静默错误", 2)
    mutation_rows = [
        ("M1 引用偏移", "=B12*C12 → =B11*C12", "引用到相邻但合法单元格"),
        ("M2 范围边界", "SUM(D2:D20) → SUM(D2:D19)", "漏一项或多一项"),
        ("M3 运算符", "A/B → A*B", "仍返回正常数值"),
        ("M4 函数", "AVERAGE → SUM", "函数语义改变"),
        ("M5 绝对引用", "$B$2 → B2 或 B$2", "复制后逐步漂移"),
        ("M6 复制偏移", "整段复制时偏移一列/一行", "形成局部结构异常"),
    ]
    add_table(doc, ["类型", "示例", "静默性"], mutation_rows, [1900, 3600, 3860], zebra=True)

    add_heading(doc, "8.3 传播深度分层", 2)
    add_table(doc, ["分层", "从源错误到可观察汇总的最短距离", "用途"], [
        ("浅层", "1—2条依赖边", "验证基本定位能力"),
        ("中层", "3—5条依赖边", "观察症状扩散后性能"),
        ("深层", "6条及以上", "检验根因与下游症状区分能力"),
    ], [1700, 4100, 3560], zebra=True)
    add_body(doc, "生成脚本先在正确工作簿上建立依赖图，再按目标深度选择可变异单元格。样本量设置为最低 800 个有效实例、目标 1200—2000 个；三种深度和六种错误类型尽量平衡。数量不是越多越好，标签验证优先于规模。")

    add_heading(doc, "8.4 有效错误的自动验收", 2)
    add_list(doc, [
        "带错工作簿可被 Excel 或 LibreOffice 正常重算，且不出现显式错误码。",
        "至少一个指定汇总输出发生数值变化；源错误确实传播到下游。",
        "至少两个下游公式受影响，避免把普通局部错当作传播错误。",
        "恢复正确公式后输出回到正确版本；错误标签与修复标签一致。",
        "错误公式与正确公式不同，且变化符合预先声明的六类变异规则。",
        "记录生成器版本、随机种子和文件哈希，保证样本可复现。",
    ])
    add_heading(doc, "8.5 数据划分防泄漏", 2)
    add_body(doc, "不能把同一个模板的不同变体随机分到训练和测试，否则算法可能记住结构。主划分按模板族隔离：开发集用于确定权重和阈值，测试集包含未见模板族；进一步做 leave-one-template-family-out 交叉验证。所有基线使用同一划分、同一候选单元格集合和同一超参数选择规则。")

    add_heading(doc, "9. 本地代码、仿真与实验环境", 1)
    add_callout(doc, "实验原则", "代码、仿真和本地测试不是“加分项”，而是论文成立的必要条件。每个数字必须能由脚本重新生成；每张表都要能追溯到结果CSV和配置文件。", fill=PALE_GOLD, title_color=GOLD)

    add_heading(doc, "9.1 建议环境", 2)
    env_rows = [
        ("操作系统", "Windows 10/11", "便于调用本机 Excel 重算"),
        ("语言", "Python 3.11+", "主算法与实验脚本"),
        ("工作簿读写", "openpyxl", "解析xlsx；注意它本身不会计算公式"),
        ("依赖图", "networkx", "DAG、可达性、深度、子图"),
        ("重算引擎A", "Excel COM + pywin32", "作为Windows主重算引擎"),
        ("重算引擎B", "LibreOffice headless", "跨引擎复核与无Excel备选"),
        ("统计与绘图", "pandas / scipy / seaborn", "结果表、置信区间、图形"),
        ("硬件", "普通笔记本，16GB内存，CPU", "不依赖GPU，便于答辩复现"),
    ]
    add_table(doc, ["项目", "选择", "说明"], env_rows, [1900, 3000, 4460], zebra=True)
    add_body(doc, "重算必须特别处理：openpyxl 只能读取或写入公式，不能像 Excel 那样计算新值。基准验证优先使用 Excel COM 批量重算，并随机抽取样本用 LibreOffice 复核；若两者结果不一致，记录为引擎差异并排除出主实验，而不是强行归入正确样本。")

    add_heading(doc, "9.2 一键实验入口", 2)
    add_body(doc, "最终提供统一命令，例如 python run.py workbook.xlsx 生成单文件诊断；python scripts/run_experiments.py --config configs/main.yaml 执行批量实验。脚本自动记录代码版本、数据版本、参数、随机种子、开始/结束时间和输出路径。")
    add_flow_table(doc, ["生成正确\n工作簿", "注入并验证\n静默错误", "批量运行\n全部方法", "计算指标\n统计检验", "生成论文\n图表与案例"])
    add_body(doc, "图2 本地仿真实验流水线。每一步失败都应中止并写入日志，防止错误数据进入最终结果。", size=9, color=MUTED, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    add_heading(doc, "10. 完整实验矩阵", 1)
    add_body(doc, "下表是必须实际运行的实验，不是论文写作提纲。优先顺序为 E0→E1→E2→E5→E7；时间不足时可以缩减扩展实验，但不能跳过工程正确性、主比较、消融和真实案例。")
    experiment_rows = [
        ("E0 工程单测", "解析器、地址平移、建图、局部重算", "pytest报告；关键模块≥80%覆盖", "全部关键测试通过"),
        ("E1 基准有效性", "800—2000个变异实例", "有效率、深度/类型分布、双引擎一致率", "标签抽查无系统性错误"),
        ("E2 主定位", "测试集所有有效错误", "Top-1/3/5、MRR、EXAM", "完整GIR优于最强无真值基线"),
        ("E3 修复候选", "每个源单元格K个候选", "Coverage@K、Repair Top-k、编辑距离", "真实修复高比例进入候选集"),
        ("E4 深度分层", "浅/中/深传播", "各层Top-k与相对下降", "解释性能随深度变化"),
        ("E5 消融", "移除各评分分量", "ΔTop-1、ΔMRR、置信区间", "证明每个核心组件价值"),
        ("E6 鲁棒/误报", "干净表、不同窗口/权重/K", "误报率、敏感性曲线", "阈值不过度依赖单设置"),
        ("E7 真实表格", "Enron错误语料可用样本+自建案例", "Top-k、案例解释、失败类型", "不是只对生成数据有效"),
        ("E8 性能", "100/500/1000/5000公式规模", "中位、P95时间与峰值内存", "常用规模可现场演示"),
    ]
    add_table(doc, ["实验", "输入/对象", "输出证据", "过关条件"], experiment_rows,
              [1450, 2650, 2900, 2360], font_size=8.15, zebra=True)

    add_heading(doc, "10.1 E0：工程正确性测试", 2)
    add_list(doc, [
        "为公式解析准备至少 60 个手工真值用例，覆盖相对/绝对/混合引用、跨表引用、范围、嵌套函数和非法公式。",
        "为依赖图准备 15 个小型工作簿，人工画出边与最短传播深度，逐一比对程序输出。",
        "为候选平移准备行、列、块三类用例，确保生成公式不会越界或改变绝对引用。",
        "为局部重算准备 20 个微型表格，与 Excel 完整重算结果比对，误差使用明确浮点容差。",
    ])

    add_heading(doc, "10.2 E1：基准有效性审计", 2)
    add_body(doc, "脚本生成后先自动过滤，再随机抽取至少 10% 的样本人工复核。抽样需覆盖全部模板族、错误类型和深度层。人工记录“标签正确/错误、是否静默、传播是否真实、修复是否恢复”四项；如果任何一类错误的标签错误率超过 5%，暂停主实验，修复生成器并重新生成该类数据。")

    add_heading(doc, "10.3 E2：主定位对比", 2)
    add_body(doc, "对每个测试工作簿，所有方法输出同一候选公式集合的排名。主单位是“工作簿”，不是单元格：先在每个工作簿计算指标，再跨工作簿汇总，避免大表格支配平均值。报告总体结果、按模板族结果、按错误类型结果和按传播深度结果。")

    add_heading(doc, "10.4 E3：候选修复", 2)
    add_body(doc, "先测候选覆盖率，判断真实公式是否进入 C(v)；再在覆盖样本上测排序效果。必须同时报告“全体样本修复率”和“候选已覆盖条件下的修复率”，否则可能把候选生成失败隐藏在排序指标之后。语义等价公式可通过多组输入重算确认，不只做字符串完全匹配。")

    add_heading(doc, "10.5 E4：传播深度实验", 2)
    add_body(doc, "每种错误类型在浅、中、深三层尽量有相近样本量。绘制传播深度—Top-1、传播深度—MRR 曲线，并统计错误源、受影响中间节点和最终汇总节点在各方法排名中的位置。该实验是本项目区别于普通公式异常检测的关键证据。")

    add_heading(doc, "10.6 E5：消融实验", 2)
    ablation_rows = [
        ("Full", "完整 FormulaGuard", "主结果"),
        ("−Formula", "移除公式族不一致", "检验公式模式贡献"),
        ("−Graph", "移除局部图模式", "检验依赖结构贡献"),
        ("−Behavior", "移除行为/扰动一致性", "检验无真值行为信号"),
        ("−Influence", "I(v)=1", "检验传播影响因子"),
        ("−Intervention", "只对原表计算异常，不替换候选", "检验“图干预”是否真正必要"),
        ("K sensitivity", "K=3/5/10/20", "覆盖率与耗时权衡"),
    ]
    add_table(doc, ["版本", "改变", "回答的问题"], ablation_rows, [1850, 3750, 3760], zebra=True)

    add_heading(doc, "10.7 E6：鲁棒性与干净表误报", 2)
    add_body(doc, "在未注入错误的正确工作簿上运行工具。由于排名系统总能给出第一名，必须另外设置“是否报警”阈值；用开发集选择阈值，再在干净测试集报告工作簿级误报率。对窗口大小、权重、候选数和阈值做敏感性分析，展示主结论是否稳定。")

    add_heading(doc, "10.8 E7：真实表格外部验证", 2)
    add_body(doc, "优先使用 Enron Error Corpus 中能确认错误公式和修复的样本，并按本项目支持语法筛选。筛选规则必须在看结果前固定，报告总样本数、排除原因和最终可用数。另选 3—5 个结构清晰的自建真实场景做完整案例，不把手工挑选的成功案例冒充总体统计。")

    add_heading(doc, "10.9 E8：运行时间与规模", 2)
    add_body(doc, "用不同公式数量和图密度的工作簿测量解析、建图、候选生成、干预评分和报告生成的分阶段耗时；每个规模重复至少 5 次，报告中位数与 P95。若大表过慢，先做候选预筛并缓存下游子图，而不是削减测试或只报最快一次。")

    add_heading(doc, "11. 对比方法与公平性", 1)
    baseline_rows = [
        ("Random", "随机打乱候选公式", "下界", "无"),
        ("Graph-only", "图影响/局部图异常排序", "验证图结构本身", "无"),
        ("Pattern-only", "邻域公式模板离群", "强简单基线", "无"),
        ("Excel提示", "Excel内置不一致公式提示", "用户熟悉的现实参照", "无"),
        ("ExceLint-like", "信息论/空间公式异常的可复现适配", "强无真值基线", "无"),
        ("WARDER-like", "公式族聚类与候选修复适配", "公式族方法参照", "无"),
        ("SEDMR-like", "按公开蜕变关系思想实现的检测适配", "行为一致性参照", "无"),
        ("SFL-Oracle", "用已知受影响输出构造失败测试", "有真值上界", "需要"),
        ("FormulaGuard", "公式+图+行为+候选干预", "本方法", "无"),
    ]
    add_table(doc, ["方法", "实现", "角色", "需要输出真值"], baseline_rows,
              [1650, 4050, 2350, 1310], font_size=8.4, zebra=True)
    add_body(doc, "若某论文没有公开代码，必须在名称后加“‑like”或“reimplementation”，列出依据的论文段落、无法复现的细节和本项目采用的替代设置；不得把自写近似实现直接称作原作者系统。ExceLint 旧版代码可能存在构建兼容问题，可先尝试原实现，再保留透明的适配版本。[6][13]")
    add_heading(doc, "11.1 公平比较协议", 2)
    add_list(doc, [
        "所有无真值方法只读取带错工作簿；任何标签文件均通过独立评测进程隔离。",
        "同一开发集选择超参数，同一测试集一次性出最终结果；不得为 FormulaGuard 单独使用更多信息。",
        "报告候选空间差异；如果某基线只能检查特定单元格，应同时给出共同候选集结果和原生设置结果。",
        "SFL-Oracle 单列为有真值上界，不用它证明无真值方法“公平胜出”。",
        "公开运行命令、配置、种子和失败日志；超时或崩溃也计入结果。",
    ])

    add_heading(doc, "12. 指标、统计与图表", 1)
    metric_rows = [
        ("Top-k Accuracy", "源错误是否出现在前k名", "k=1/3/5，最直观"),
        ("MRR", "1 / 源错误名次", "强调前排质量"),
        ("EXAM", "检查到源错误前需查看的公式比例", "跨大小表格可比"),
        ("Candidate Coverage@K", "正确修复是否在候选集", "分离生成与排序"),
        ("Repair Top-k", "正确/语义等价修复是否前k", "修复实用性"),
        ("Clean FPR", "干净工作簿被报警的比例", "避免工具逢表必报"),
        ("Runtime / Memory", "中位、P95、峰值内存", "现场可用性"),
    ]
    add_table(doc, ["指标", "定义", "作用"], metric_rows, [2100, 4300, 2960], zebra=True)
    add_body(doc, "统计上以工作簿为独立单位，使用分层 bootstrap（建议 1000 次）给出 95% 置信区间；对同一工作簿上两种方法的排名指标做配对 Wilcoxon 检验，并同时报告效应量。显著性不能替代实际提升：即使 p 值很小，若 Top-1 只提升 1 个百分点，也不应宣传为重大突破。")
    add_heading(doc, "12.1 最终论文建议图表", 2)
    add_list(doc, [
        "表：总体主结果（无真值组）和有真值上界分开。",
        "图：传播深度 vs Top-1/MRR 折线，带 95% 置信区间。",
        "图：不同错误类型的 Top-5 分组柱状图。",
        "表：消融结果与相对下降。",
        "图：候选数 K 的覆盖率—运行时间双轴或两图。",
        "案例图：一个真实错误的依赖子图、候选替换和能量变化。",
        "表：失败案例分类，不隐藏方法弱点。",
    ])

    doc.add_page_break()
    add_heading(doc, "13. 软件实现计划与仓库结构", 1)
    repo_rows = [
        ("src/formula_parser.py", "AST解析、A1/R1C1归一化"),
        ("src/family.py", "公式指纹、邻域和公式族"),
        ("src/dependency_graph.py", "依赖边、深度、下游子图"),
        ("src/candidates.py", "候选公式生成与编辑成本"),
        ("src/energy.py", "三类不一致能量"),
        ("src/gir.py", "干预重算、GIR排序和解释"),
        ("src/recalculate.py", "Excel COM/LibreOffice重算适配"),
        ("bench/generators/", "8—10个工作簿模板"),
        ("bench/mutators.py", "六类错误注入和标签记录"),
        ("baselines/", "全部对比方法，接口统一"),
        ("scripts/run_experiments.py", "读取YAML配置并批量执行"),
        ("tests/", "解析、建图、平移、重算、评分单测"),
        ("results/", "CSV、图、日志；按实验编号组织"),
        ("report/", "论文图表索引、案例和答辩材料"),
    ]
    add_table(doc, ["路径", "职责"], repo_rows, [3350, 6010], font_size=8.8, zebra=True)

    add_heading(doc, "13.1 模块接口", 2)
    add_body(doc, "所有定位器实现统一接口 rank(workbook_path) → [{cell, score, repair, evidence}]。基准评测器只读取结果 JSON 和隐藏标签，防止算法误读真值。每次实验生成 manifest.json，记录 Git 提交号、配置哈希、数据哈希和环境版本。")
    add_heading(doc, "13.2 最小可用版本（MVP）", 2)
    add_list(doc, [
        "能读取单工作簿，解析支持范围内的公式并构建依赖图。",
        "能识别公式族、生成邻居平移候选并计算简化 GIR。",
        "能在 20 个手工微型表格上输出 Top-5 和影响路径。",
        "能自动注入 M1/M2 两类错误并生成隐藏标签。",
        "能一键运行 Random、Pattern-only、Graph-only 和 FormulaGuard。",
    ])
    add_body(doc, "MVP必须在第二周末完成。后续所有高级特征都建立在它之上；如果时间紧张，优先提高基准质量和实验可信度，不追求堆叠更多公式函数或模型。")

    add_heading(doc, "14. 进度安排（2026.08.10—09.15）", 1)
    schedule_rows = [
        ("08.10—08.12", "锁定问题、范围和文献矩阵", "研究问题、原创边界、20个微型表格", "范围不再随意扩张"),
        ("08.13—08.18", "解析、归一化、依赖图、单测", "E0初版；可画依赖路径", "核心单测通过"),
        ("08.19—08.23", "公式族、候选、简化GIR", "MVP命令行工具", "20例Top-5可检查"),
        ("08.24—08.28", "模板生成器与六类变异", "PropagationBench v0.1", "≥300有效实例"),
        ("08.29—09.02", "基线、批量实验和消融", "E1/E2/E5初版结果", "能形成主结果表"),
        ("09.03—09.06", "扩充样本、深度/修复/性能", "E3/E4/E8结果", "结果可复现"),
        ("09.07—09.09", "真实表格、失败案例、误报", "E6/E7与2个完整案例", "不只在生成数据有效"),
        ("09.10—09.12", "论文、图表、系统演示视频", "完整初稿与演示", "指导教师完成首轮审阅"),
        ("09.13—09.14", "复现实验、合规、答辩演练", "最终包、AI披露、备份", "冻结代码和结果"),
        ("09.15", "仅做提交与应急", "按官方系统完成提交", "24:00前完成并留凭证"),
    ]
    add_table(doc, ["日期", "任务", "产物", "阶段门"], schedule_rows,
              [1500, 2750, 2800, 2310], font_size=8.05, zebra=True)
    add_callout(doc, "时间管理红线", "9月12日后禁止加入新算法模块；9月13日起只做复现、修错、排版和演练。9月15日不安排关键实验，以免提交前出现不可恢复问题。", fill=PALE_RED, title_color=RED)

    add_heading(doc, "15. 人员分工与每周工作节奏", 1)
    add_body(doc, "以下按 1—3 名学生设计，可根据真实团队填写。即使分工，每名学生都必须理解完整问题和主实验；核心代码、数据和论文不能成为某一个人的“黑箱”。")
    role_rows = [
        ("学生A/负责人", "依赖图、GIR、实验总控、论文主线", "能从零解释算法与主结果"),
        ("学生B（可选）", "基准生成、错误注入、重算验证", "能说明标签为何可信"),
        ("学生C（可选）", "基线复现、统计图表、真实案例", "能说明比较为何公平"),
        ("指导教师", "范围把关、学术规范、方法质疑、答辩训练", "不代写核心代码与结论"),
    ]
    add_table(doc, ["角色", "职责", "答辩要求"], role_rows, [1900, 4150, 3310], zebra=True)
    add_heading(doc, "15.1 建议日节奏", 2)
    add_list(doc, [
        "每天开始：写下唯一一个可验证目标，例如“让范围边界变异通过20个测试”。",
        "每天结束：提交代码、保存日志、记录失败原因和第二天动作，不用口头记忆代替记录。",
        "每两天：从空环境运行一次主命令，防止只有开发者电脑能运行。",
        "每周：向指导教师做一次10分钟演示，重点讲证据和失败，不只展示界面。",
    ])

    add_heading(doc, "16. 风险、止损与备选方案", 1)
    risk_rows = [
        ("公式解析范围过大", "高", "高", "冻结支持语法；不支持项明确记录并跳过", "MVP两天内仍不稳定"),
        ("Excel批量重算不稳定", "中", "高", "进程隔离、超时、备份副本；LibreOffice备选", "连续>5%文件失败"),
        ("候选覆盖率低", "中", "高", "先强化邻域/块模板，再考虑AST编辑", "Coverage@10<70%"),
        ("GIR不优于模式基线", "中", "高", "查深度分层和能量分量；收缩主张", "两轮修正仍无稳定提升"),
        ("真实语料难复现", "高", "中", "固定筛选规则；以案例分析补充，不伪造总体", "可用真实样本<15"),
        ("运行速度过慢", "中", "中", "候选预筛、缓存、局部下游重算", "1000公式P95>60秒"),
        ("时间被论文挤压", "中", "高", "9月7日起边实验边写；12日冻结功能", "9月10日仍无完整初稿"),
        ("AI使用不合规", "低", "高", "事先获教师同意；保留工具/版本/提示记录", "无法还原使用记录"),
    ]
    add_table(doc, ["风险", "概率", "影响", "缓解", "触发止损"], risk_rows,
              [1700, 720, 720, 3900, 2320], font_size=7.85, zebra=True)
    add_body(doc, "止损并不等于放弃项目，而是把论文主张缩到已被证据支持的范围。例如，如果修复排序不稳定，仍可把主要成果限定为根因定位和候选覆盖；如果真实样本过少，就把真实部分定位为案例研究，不能把小样本包装为大规模验证。")

    add_heading(doc, "17. 银奖目标的内部验收门槛", 1)
    add_body(doc, "下列数值是团队在实验前设定的“竞争力目标”，用于判断是否需要改进或收缩主张，不是奖项保证，也不是当前结果。最终论文必须报告真实结果，即使未达目标。")
    gate_rows = [
        ("G1 工程", "关键单测全部通过；主实验可一键复现", "必须"),
        ("G2 数据", "≥800个有效变异；标签抽查无系统性错误", "必须"),
        ("G3 主结果", "Top-1相对最强无真值基线提升≥10个百分点，且95%CI稳定", "强目标"),
        ("G4 实用性", "Top-5≥80%；用户最多检查5个公式", "强目标"),
        ("G5 修复", "Candidate Coverage@10≥85%；Repair Top-1≥50%", "目标"),
        ("G6 真实", "真实可用错误样本Top-5约≥60%，并给出失败分析", "目标"),
        ("G7 误报", "干净工作簿报警率≤15%（阈值由开发集固定）", "目标"),
        ("G8 性能", "≤1000公式工作簿P95运行时间≤30秒", "目标"),
        ("G9 展示", "3分钟演示可完整重现一个定位案例", "必须"),
        ("G10 学术", "文献、代码来源、AI使用和失败结果全部披露", "必须"),
    ]
    add_table(doc, ["门槛", "判定", "性质"], gate_rows, [1650, 6410, 1300], zebra=True)
    add_callout(doc, "真正影响奖项的证据", "优先级依次是：标签可信 → 公平基线 → 稳定提升 → 真实案例 → 可解释演示 → 排版包装。漂亮界面不能补救无效实验，复杂公式也不能替代清楚的问题和证据。", fill=PALE_GREEN, title_color=GREEN)

    doc.add_page_break()
    add_heading(doc, "18. 论文、演示与最终交付物", 1)
    add_heading(doc, "18.1 论文建议结构", 2)
    paper_rows = [
        ("摘要", "问题、方法、数据、最关键结果和限制；结果出来后最后写"),
        ("1 引言", "静默错误案例、无真值难点、贡献列表"),
        ("2 相关工作", "按调试、异常、修复、蜕变测试分类，界定差异"),
        ("3 问题定义", "工作簿图、源错误、静默性、oracle-free输入输出"),
        ("4 方法", "归一化、公式族、图、能量、候选、GIR、复杂度"),
        ("5 基准", "模板、变异、深度、验证、划分和统计"),
        ("6 实验", "RQ、基线、指标、主结果、消融、真实验证、性能"),
        ("7 讨论", "为何有效、失败案例、威胁、适用边界"),
        ("8 结论", "回答RQ，不加入未验证的新主张"),
    ]
    add_table(doc, ["部分", "必须回答"], paper_rows, [1900, 7460], zebra=True)

    add_heading(doc, "18.2 最终交付清单", 2)
    add_list(doc, [
        "研究论文 PDF 与可编辑源文件。",
        "FormulaGuard 源代码、依赖锁定文件、README 和开源许可证说明。",
        "PropagationBench 生成脚本、可公开样本、隐藏标签评测脚本和数据说明。",
        "全部实验配置、原始 CSV、统计脚本、图表生成脚本和结果索引。",
        "3—5分钟系统演示视频：加载表格、输出Top-5、查看影响路径、验证修复。",
        "答辩PPT、1页项目速览、失败案例清单和常见问题答案。",
        "第三方代码/数据来源清单、AI辅助使用说明和必要的对话记录。",
        "离线备份与一键运行包，确保现场网络不稳定也能演示。",
    ])

    add_heading(doc, "18.3 答辩必须能回答的问题", 2)
    add_list(doc, [
        "为什么不是简单的 ExceLint + 依赖图？请用消融和干预定义回答。",
        "算法没有正确输出，怎样知道某个替换更好？请解释不一致能量与真值隔离。",
        "为什么你定位的是源错误，而不是影响最大的上游节点？请展示反例和候选替换结果。",
        "基准是不是为自己的方法量身定做？请解释模板隔离、真实表格和公平基线。",
        "哪些样本失败？失败是否集中在唯一公式、复杂函数或跨表结构？",
        "代码中你本人完成了什么？能否现场修改一个参数并重新运行？",
    ])

    add_heading(doc, "19. 学术诚信与AI使用", 1)
    add_body(doc, "竞赛官方对人工智能工具的使用有专门要求，应以官方最新规则和指导教师意见为准。[16] 本项目可以使用 AI 帮助查找关键词、解释文献、检查代码、润色表达和生成测试想法，但研究问题、算法决策、代码验证、实验执行、数据判断和结论必须由学生理解并负责。")
    ai_rows = [
        ("使用前", "取得指导教师同意；记录允许的工具范围；确认官方当年规则"),
        ("使用中", "记录工具名称、版本、日期、用途、关键提示、采用/未采用内容"),
        ("代码", "逐行理解后再采用；标出第三方或AI辅助部分；用测试证明正确"),
        ("文献", "AI给出的题名、作者、年份必须回到论文或官方页面核实"),
        ("论文", "不得让AI伪造实验、引用、图表或结论；保留写作迭代记录"),
        ("提交", "按规则提交披露表和需要的对话记录；学生能独立答辩"),
    ]
    add_table(doc, ["阶段", "要求"], ai_rows, [1900, 7460], zebra=True)
    add_callout(doc, "本计划本身的披露", "本研究计划书使用了 OpenAI Codex 辅助结构化、措辞与版式生成。参赛团队应由学生和指导教师逐条审核、修改并确认，后续按竞赛规则如实披露，不应把本文件直接当作学生独立完成的最终论文。", fill=PALE_GOLD, title_color=GOLD)

    add_heading(doc, "20. 参考资料与官方链接", 1)
    references = [
        "[1] Erwig, M. et al. GoalDebug: A Spreadsheet Debugger for End Users. ICSE, 2007. https://web.engr.oregonstate.edu/~erwig/papers/GoalDebug_ICSE07.pdf",
        "[2] Abraham, R.; Erwig, M. Spreadsheet Debugging. 2008. https://arxiv.org/abs/0801.4280",
        "[3] Hofer, B. et al. On the Empirical Evaluation of Fault Localization Techniques for Spreadsheets. 2013. https://researchportal.ulisboa.pt/pt/publications/on-the-empirical-evaluation-of-fault-localization-techniques-for-/",
        "[4] Hofer, B. et al. The FaultySheet Detective: When Smells Meet Fault Localization. 2014. https://gzoltar.com/pub/18.pdf",
        "[5] MUSSCO: A Mutation-Based Spreadsheet Formula Repair Approach. 2015. https://dl.ifip.org/IFIP-LNCS-9447/hal-01470160",
        "[6] Barowy, D. W. et al. ExceLint: Automatically Finding Spreadsheet Formula Errors. OOPSLA, 2018. https://www.microsoft.com/en-us/research/publication/excelint-automatically-finding-spreadsheet-formula-errors/",
        "[7] Metric-based Fault Prediction for Spreadsheets, project resources. https://spreadsheet-research.github.io/Metric-based-Fault-Prediction-for-Spreadsheets/",
        "[8] WARDER: spreadsheet formula error detection/repair through formula structure and clustering. 2020. https://www.sciencedirect.com/science/article/pii/S0164121220300935",
        "[9] Bavishi, R. et al. LaMirage: Neurosymbolic Repair for Low-Code Formula Languages. 2022. https://www.microsoft.com/en-us/research/publication/neurosymbolic-repair-for-low-code-formula-languages/",
        "[10] FLAME: A Small Language Model for Spreadsheet Formulas. 2023. https://arxiv.org/abs/2301.13779",
        "[11] FoRepBench: A Benchmark for Context-Aware Spreadsheet Formula Repair. 2025. https://kdd-eval-workshop.github.io/genai-evaluation-kdd2025/assets/papers/Submission%2033.pdf",
        "[12] SEDMR: Spreadsheet Error Detection Based on Metamorphic Relations. 2026. https://www.sciencedirect.com/science/article/abs/pii/S0950584926000637",
        "[13] ExceLint source repository. https://github.com/ExceLint/ExceLint",
        "[14] Enron Error Corpus. https://spreadsheets.sai.tugraz.at/index.php/corpora-for-benchmarking/enron-error-corpus/",
        "[15] 丘成桐中学科学奖：2026日程、评审标准与参赛规则。https://www.yau-awards.com/page-schedule.html ; https://www.yau-awards.com/page-criteria.html ; https://www.yau-awards.com/page-rule.html",
        "[16] 丘成桐中学科学奖：人工智能工具使用规则。https://www.yau-awards.com/en/show-86-59.html",
    ]
    for ref in references:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(ref)
        set_run_font(r, size=8.55, color=NAVY)
    add_body(doc, "注：正式论文写作时，应下载并阅读原文，补齐作者、期刊/会议、卷期页码和 DOI；不得只根据网页摘要或AI转述引用。官方日程和规则可能更新，提交前再次核对官网。", size=9, color=RED, after=8)

    doc.add_page_break()
    add_heading(doc, "附录A：单个错误实例记录模板", 1)
    record_rows = [
        ("instance_id", "唯一编号，例如 budget_03_M2_d4_seed17"),
        ("clean_workbook", "正确工作簿文件及SHA-256"),
        ("mutant_workbook", "带错工作簿文件及SHA-256"),
        ("source_cell", "源错误工作表与地址"),
        ("original_formula", "正确公式，仅评测端可见"),
        ("mutated_formula", "错误公式"),
        ("mutation_type", "M1—M6"),
        ("propagation_depth", "到指定汇总输出的最短路径"),
        ("affected_cells", "受影响下游列表或数量"),
        ("recalc_engine", "Excel/LibreOffice版本"),
        ("validation", "静默、传播、可恢复、双引擎一致"),
        ("seed / generator", "随机种子与生成器版本"),
    ]
    add_table(doc, ["字段", "内容"], record_rows, [2500, 6860], zebra=True)

    add_heading(doc, "附录B：每次主实验运行清单", 1)
    checklist = [
        "□ 工作区干净，代码已提交，记录当前提交号。",
        "□ 数据 manifest 与哈希已保存，隐藏标签未进入算法路径。",
        "□ 配置文件只从开发集确定，未查看测试结果后修改。",
        "□ 随机种子、环境版本、Excel/LibreOffice版本已记录。",
        "□ 全部基线和FormulaGuard使用相同测试列表。",
        "□ 崩溃、超时和跳过样本均写入失败日志。",
        "□ 原始结果CSV只追加不手工修改；清洗由脚本完成。",
        "□ 统计图表可以由一个命令从原始CSV重新生成。",
        "□ 抽查至少5个样本的排名、候选和影响路径。",
        "□ 备份结果、日志、配置和生成图表。",
    ]
    for item in checklist:
        add_body(doc, item, size=10.4, after=5, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "附录C：每日研究日志模板", 1)
    log_rows = [
        ("日期 / 用时", ""),
        ("今天唯一目标", ""),
        ("执行命令与配置", ""),
        ("得到的证据", ""),
        ("失败与原因", ""),
        ("是否改变原假设", ""),
        ("AI/外部帮助记录", "工具、版本、用途、采用内容"),
        ("明天第一步", ""),
    ]
    table = add_table(doc, ["项目", "记录"], log_rows, [2500, 6860], zebra=False)
    for row in table.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(16)

    add_spacer(doc, 10)
    add_callout(doc, "最终提醒", "先把最小系统跑通，再扩展；先保证标签正确，再追求样本规模；先做公平对比，再讨论创新；先让自己能讲明白，再追求术语复杂。做到这些，才是这份计划真正接近银奖竞争力的地方。", fill=BLUE_GRAY, title_color=NAVY)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
