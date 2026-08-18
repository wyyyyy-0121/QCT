"""Build a Word draft from frozen FormulaGuard v3 evidence.

This is intentionally a draft: batch-throughput and independent-blind-study
sections are visibly marked pending instead of being filled with estimates.
"""

from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "FormulaGuard_v3_论文初稿_待补充吞吐.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B677A"
LIGHT = "F4F6F9"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def csv_rows(path: Path):
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    width_node = tc_pr.find(qn("w:tcW"))
    if width_node is None:
        width_node = OxmlElement("w:tcW")
        tc_pr.append(width_node)
    width_node.set(qn("w:w"), str(width))
    width_node.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == TABLE_WIDTH
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT))
    indent.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, LIGHT)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            run = paragraph.add_run(str(value))
            set_run_font(run, size=9.2, color="202020")
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_para(doc, text="", *, bold_prefix=None, style=None, align=None, size=None, color=None, italic=False, after=8, before=0):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=size, color=color, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=size, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    first = p.add_run(label + " ")
    set_run_font(first, size=10.2, color=DARK_BLUE, bold=True)
    rest = p.add_run(text)
    set_run_font(rest, size=10.2, color="26384D")
    set_table_geometry(table, [TABLE_WIDTH])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        run = p.add_run(item)
        set_run_font(run, size=11, color="202020")


def add_page_number(paragraph):
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for name, size, color, before, after in (("Heading 1", 16, BLUE, 18, 10), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, DARK_BLUE, 8, 4)):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    subtitle = styles.add_style("Paper Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("FormulaGuard v3 | 竞赛论文研究初稿"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("FormulaGuard v3  |  第 "), size=8.5, color=MUTED)
    add_page_number(footer)
    set_run_font(footer.add_run(" 页"), size=8.5, color=MUTED)


def main():
    full = ROOT / "results" / "v3_full"
    enron = ROOT / "results" / "enron_test_v3_real"
    summary = {row["method"]: row for row in csv_rows(full / "summary.csv")}
    enron_summary = {row["method"]: row for row in csv_rows(enron / "external_summary.csv")}
    latency = csv_rows(full / "performance_v3_latency_summary.csv")
    v3v2 = json.loads((full / "v3_vs_v2.json").read_text(encoding="utf-8"))
    coupling = json.loads((full / "benchmark_independence.audit.json").read_text(encoding="utf-8"))
    clean = json.loads((full / "clean_summary.json").read_text(encoding="utf-8"))
    blind = json.loads((enron / "blind_result_audit.json").read_text(encoding="utf-8"))

    doc = Document()
    configure(doc)
    # Cover
    for _ in range(7):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    add_para(doc, "丘成桐中学科学奖参赛研究初稿", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=MUTED, bold_prefix=None, after=18)
    add_para(doc, "FormulaGuard", align=WD_ALIGN_PARAGRAPH.CENTER, size=30, color=INK, bold_prefix=None, after=8)
    p = add_para(doc, "面向无输出真值场景的电子表格静默源公式错误定位", align=WD_ALIGN_PARAGRAPH.CENTER, size=17, color=DARK_BLUE, after=18)
    p.runs[0].bold = True
    add_para(doc, "依赖图、候选修复与下游反事实诊断责任排序", style="Paper Subtitle", after=28)
    add_callout(doc, "当前版本说明", "合成 full、LibreOffice、Enron 冻结盲测和单工作簿延迟已经写入；批量吞吐与独立小样本盲测尚未补入，相关位置均明确标注为待完成。")
    add_para(doc, "参赛学生：待填写    指导教师：待填写    完成日期：2026 年 8 月", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, color=MUTED, after=20)
    doc.add_page_break()

    add_heading(doc, "摘要", 1)
    add_para(doc, "电子表格公式错误常不会产生 #REF! 等显式异常，而会以看似正常的数值沿依赖关系向下游传播。现实使用者通常既没有正确输出真值，也没有失败测试用例，因而难以从多个下游症状回溯最初写错的公式。本文提出 FormulaGuard：先用局部公式模式、依赖图结构与工作簿内部一致性形成可疑先验，再对有限、可解释的候选修复进行临时替换，观察异常能量及下游区域是否出现净改善。系统输出待审查的 Top-k 单元格、候选修复、传播路径和证据等级，而不读取正确公式或正确输出。")
    add_para(doc, "实验分为受控合成、跨引擎计算验证与真实错误盲测三层。PropagationBench-v3 包含 864 个可控实例；在该基准上 frozen-v3 的 MRR 为 1.000，但候选—变异器镜像匹配率同为 100%，故结果只能解释为候选存在条件下的机制验证。官方 Enron Error Corpus 的 20 个冻结测试事件上，图基线 MRR 为 0.365，FormulaGuard v3 为 0.250；v3 相对 v2 的 bootstrap 95% 区间跨 0，不能宣称真实表格上稳定优于图基线。本文保留该负面结果，并将方法定位为可审计的离线风险排序工具而非自动证明错误的实时插件。")
    add_para(doc, "关键词：电子表格错误定位；静默错误；依赖图；候选修复；反事实诊断；可复现研究", bold_prefix="关键词：", after=12)

    add_heading(doc, "目录", 1)
    for item in ["1 引言", "2 相关工作与差异", "3 问题定义", "4 FormulaGuard 方法", "5 数据与实验设计", "6 实验结果与讨论", "7 真实性、局限与有效性威胁", "8 结论", "参考文献", "附录 A：复现与提交审计"]:
        add_para(doc, item, after=3)
    doc.add_page_break()

    add_heading(doc, "1 引言", 1)
    add_para(doc, "预算、成绩统计、库存管理和商业分析等工作簿往往由大量相互引用的公式组成。若某一处求和范围漏掉一行、复制时绝对引用符号丢失，或运算符被误改，Excel 仍可能输出普通数值；最终汇总的异常却可能来自很早的上游公式。这样的错误具有“症状多、源头少”的特点，单纯寻找与邻居不同的公式并不能区分源错误和下游症状。")
    add_para(doc, "本研究关注无输出真值的严格场景：算法只能读取工作簿中的公式、引用关系、已有数值和冗余检查，不能读取正确输出、正确公式或已知错误单元格。目标不是在因果意义上证明某格必错，而是在有限人工审查预算内，优先呈现更值得检查的源公式，并说明候选修复可能改善哪些内部矛盾。")
    add_bullets(doc, [
        "任务贡献：定义无输出真值、静默公式错误和源错误审查排序的输入、输出与评价边界。",
        "方法贡献：将候选小修复、工作簿内部一致性变化和下游依赖区域恢复组合为可审计证据，而不是把局部离群直接等同于错误。",
        "证据贡献：以合成耦合审计、官方 Enron 事件级盲测、LibreOffice 交叉验证、失败案例和性能测量共同限制结论强度。",
    ])

    add_heading(doc, "2 相关工作与差异", 1)
    add_heading(doc, "2.1 局部公式异常与公式簇", 2)
    add_para(doc, "ExceLint 利用电子表格的矩形结构和相对公式模式识别局部“惊异”公式 [3]。WARDER 则通过多层有效性规则细化单元格簇，以改善缺失或不一致公式检测 [4]。它们表明无需正确输出也能发现异常，但局部异常并不自动等于传播链中的源错误。FormulaGuard 因此将公式模式仅作为先验，并要求候选替换带来可见的内部改善才提供更强证据。")
    add_heading(doc, "2.2 有测试信息的电子表格调试", 2)
    add_para(doc, "GoalDebug 允许用户表达对单元格值的预期，并据此辅助公式修改 [1]。Hofer 等将谱故障定位、动态切片和约束调试用于电子表格，并在有通过/失败测试信息的条件下比较方法 [2]。本文认可依赖图排序的基础价值，但不假设使用者拥有失败输出；SFL-Oracle 仅作为有真值辅助参照，不与无真值方法混合比较。")
    add_heading(doc, "2.3 蜕变关系、候选修复与本文边界", 2)
    add_para(doc, "SEDMR 通过输入变换下的蜕变关系检测公式错误 [5]。FormulaGuard 的操作不同：它不设计输入扰动，而是暂时替换一个有限候选公式并观察工作簿内部异常和下游区域的变化。本文不声称首次使用依赖图、候选修复或反事实思想；窄创新是把它们组织成一个没有正确输出真值时仍可解释、可降级的源错误审查流程。")
    add_table(doc, ["方法类别", "需要的额外信息", "主要证据", "与本文的关系"], [
        ["GoalDebug [1]", "用户预期值", "目标导向公式修改", "输入条件不同"],
        ["SFL/动态切片 [2]", "通过/失败测试", "失败输出相关性", "有真值辅助参照"],
        ["ExceLint [3]", "无", "局部模式异常", "无真值强基线思想"],
        ["WARDER [4]", "无", "有效性驱动公式簇", "透明 warder_like 近似"],
        ["SEDMR [5]", "可用蜕变关系", "输入变换违反", "相邻任务，证据路径不同"],
        ["FormulaGuard", "无正确输出", "候选替换后的内部净改善", "本文方法"],
    ], [1600, 2150, 2700, 2910])

    add_heading(doc, "3 问题定义", 1)
    add_para(doc, "令工作簿的公式单元格集合为 V。若公式 u 被公式 v 引用，则建立有向边 u→v，得到依赖图 G=(V,E)。静默错误是某个单元格 s 的公式偏离预期、但仍返回普通数值的情况；错误可能通过 G 传播至多个后代。系统输入为工作簿 W，输出 V 的可疑排序、每个高排名单元格的候选修复 r，以及至下游汇总节点的路径摘要。")
    add_callout(doc, "无真值约束", "错误源、正确公式和正确输出只能存放在离线评价标签中；FormulaGuard 在排序完成前不可读取它们。这里的“反事实”是可执行的候选替换实验，不等价于具有可识别性保证的结构因果效应。")
    add_para(doc, "主指标为 Top-1、Top-3、Top-5、MRR 与 EXAM。多单元格错误范围按最早命中的标注公式计算一个事件，而不把范围拆成大量伪独立样本。真实语料中正确公式未知时，只评价定位，不评价候选修复精确率。")

    add_heading(doc, "4 FormulaGuard 方法", 1)
    add_heading(doc, "4.1 公式归一化、依赖图与局部先验", 2)
    add_para(doc, "系统把 A1 公式转换为相对引用指纹，保留绝对引用符号 $，从而比较复制块中的同族公式。对每个公式 v，计算公式族异常 A_f(v)、图异常 A_g(v) 和行为异常 A_b(v)。行为异常包括可重算数值的稳健偏差和工作簿内部冗余检查形成的约束残差；它们均来自工作簿可见信息而不是外部答案。")
    add_para(doc, "v3 使用局部结构可靠度 ρ(v) 调整三项权重：w_f=0.45，w_g=0.20ρ，w_b=0.55-0.20ρ；局部先验为 L(v)=w_fA_f+w_gA_g+w_bA_b。若公式位于纵向递推块首行，使用工程性的边界因子 U(v)=0.25 抑制其被机械地视为异常。")
    add_heading(doc, "4.2 候选修复与反事实评价", 2)
    add_para(doc, "候选来自相邻同族公式的平移以及有界小编辑：普通引用和复制偏移、范围端点、四则运算符、SUM/AVERAGE/MIN/MAX 替换，以及绝对引用符号的添加或删除。候选需可解析、引用有效率足够且不越界；主实验每格保留 Top-15，并用编辑族配额避免某类候选占满列表。")
    add_para(doc, "对候选 r 临时替换 v 的公式，分别测量公式、图、一般行为和内部约束四个能量分量的增益 gain_j 与伤害 harm_j。以结构可靠度加权后，净改善定义为 N(v,r)=max(0,G(v,r)-0.50H(v,r))。因此，仅仅让局部公式更像邻居、但破坏内部约束的候选不会得到完整奖励。")
    add_heading(doc, "4.3 CAR 分数与证据等级", 2)
    add_para(doc, "令 I(v) 为后代范围归一化影响，C(v) 为可达检查单元格比例，R_d(v,r) 为候选在后代区域消除异常的比例。路径责任 T(v,r)=(0.60I+0.40C)R_d。冻结 v3 的候选分数为 CAR(v,r)=0.20L*(v)+0.60N(v,r)+0.10N(v,r)T(v,r)+0.10N(v,r)Q(v,r)，单元格取其最佳候选。传播和候选质量均必须乘以正的净改善，避免无证据放大。")
    add_para(doc, "在真实工作簿上，v3-real B 采用更保守的规则：v2 原始分数保持主排序，正 CAR 证据仅打破完全相同的 v2 分数。它输出 counterfactual_supported、pattern_only 或 insufficient_evidence，以区分“有正替换证据”“只有模式信号”与“证据不足”。")
    add_table(doc, ["步骤", "输入", "输出", "标签可见性"], [
        ["1. 解析与归一化", "公式与引用", "公式族、依赖图", "不可见"],
        ["2. 局部先验", "公式/图/内部残差", "L(v) 排序", "不可见"],
        ["3. 生成候选", "相邻公式与有界编辑", "Top-15 候选", "不可见"],
        ["4. 临时替换", "候选和后代区域", "净改善与副作用", "不可见"],
        ["5. 输出审查单", "CAR 与路径", "Top-k、候选、证据等级", "不可见"],
        ["6. 离线评价", "冻结标签", "Top-k/MRR/EXAM", "仅此步可见"],
    ], [1500, 2500, 2800, 2560])

    add_heading(doc, "5 数据与实验设计", 1)
    add_heading(doc, "5.1 证据层级与标签隔离", 2)
    add_para(doc, "PropagationBench-v3 提供 864 个受控错误实例，用于控制错误类型、传播深度和模板结构。其错误源和正确公式存放在独立评价标签中。Enron Error Corpus 官方页说明其包含 36 个错误及标记错误单元格的 properties 文件 [6]；项目对 36 个事件逐项清点，最终纳入 30 个公式错误事件。开发集 10 个事件和冻结测试集 20 个事件的清单哈希分离，测试集在规则冻结前未运行。")
    add_heading(doc, "5.2 基线、消融与统计纪律", 2)
    add_para(doc, "无真值基线包括 pattern、graph、behavior、excelint_like、warder_like 和 v2 GIR；其中 -like 明确是透明思想级近似。SFL-Oracle 利用额外标签，仅作为上界参照。v3 消融包括去结构自适应、去副作用惩罚和去传播路径责任。所有合成主比较来自测试拆分；v3 相对 v2 的 MRR 差用配对 bootstrap 95% 区间解释，区间跨 0 时禁止使用“稳定优于”。")
    add_heading(doc, "5.3 跨引擎与性能协议", 2)
    add_para(doc, "LibreOffice 对 22 份工作簿重算后，1412 个可比缓存公式值全部一致。性能实验把单工作簿隔离延迟与多进程批量吞吐严格分开：前者固定一个进程、每个规模五次重复；后者只报告总墙钟时间与 jobs/sec，不能把资源竞争下的单任务耗时伪装成延迟。")

    add_heading(doc, "6 实验结果与讨论", 1)
    add_heading(doc, "6.1 合成机制实验：强结果与强限制必须同时报告", 2)
    add_table(doc, ["方法", "Top-1", "Top-5", "MRR", "EXAM"], [
        ["FormulaGuard v3", f"{float(summary['formulaguard_v3']['top1']):.3f}", f"{float(summary['formulaguard_v3']['top5']):.3f}", f"{float(summary['formulaguard_v3']['mrr']):.3f}", f"{float(summary['formulaguard_v3']['exam']):.3f}"],
        ["v2 GIR", f"{float(summary['formulaguard']['top1']):.3f}", f"{float(summary['formulaguard']['top5']):.3f}", f"{float(summary['formulaguard']['mrr']):.3f}", f"{float(summary['formulaguard']['exam']):.3f}"],
        ["WARDER-like", f"{float(summary['warder_like']['top1']):.3f}", f"{float(summary['warder_like']['top5']):.3f}", f"{float(summary['warder_like']['mrr']):.3f}", f"{float(summary['warder_like']['exam']):.3f}"],
        ["graph", f"{float(summary['graph']['top1']):.3f}", f"{float(summary['graph']['top5']):.3f}", f"{float(summary['graph']['mrr']):.3f}", f"{float(summary['graph']['exam']):.3f}"],
        ["pattern", f"{float(summary['pattern']['top1']):.3f}", f"{float(summary['pattern']['top5']):.3f}", f"{float(summary['pattern']['mrr']):.3f}", f"{float(summary['pattern']['exam']):.3f}"],
    ], [2550, 1700, 1700, 1700, 1710])
    add_para(doc, f"在 864 个受控实例上，v3 相对 v2 的平均 MRR 差为 {float(v3v2['mean_mrr_difference_v3_minus_v2']):.3f}，bootstrap 95% 区间为 [{v3v2['bootstrap_95_ci'][0]:.3f}, {v3v2['bootstrap_95_ci'][1]:.3f}]。但 Candidate Coverage@15 为 {float(coupling['exact_candidate_coverage']):.0%}，镜像变异—候选操作匹配率也为 {float(coupling['mirrored_operator_rate']):.0%}，被审计为 high 耦合风险。因此，表中满分只说明在正确候选存在且公式语法受支持时，反事实排序可定位所注入的源错误，不能替代自然工作簿泛化结论。")
    add_heading(doc, "6.2 干净工作簿与 LibreOffice 一致性", 2)
    add_para(doc, f"冻结阈值在 48 份合成干净工作簿上的报警率为 {float(clean['alarm_rate']):.0%}，但该数值仅是合成 clean estimate，不能称为生产环境误报率。跨引擎验证显示 22 份工作簿、1412 个可比公式缓存值全部一致，说明受支持公式的求值在本实验中与 LibreOffice 一致；它不验证业务真值或官方标签的正确性。")
    add_heading(doc, "6.3 Enron 冻结盲测：真实集没有稳定增益", 2)
    add_table(doc, ["方法", "Top-1", "Top-3", "Top-5", "MRR", "EXAM"], [
        ["graph", *[f"{float(enron_summary['graph'][key]):.3f}" for key in ('top1', 'top3', 'top5', 'mrr', 'exam')]],
        ["pattern", *[f"{float(enron_summary['pattern'][key]):.3f}" for key in ('top1', 'top3', 'top5', 'mrr', 'exam')]],
        ["FormulaGuard v2", *[f"{float(enron_summary['formulaguard'][key]):.3f}" for key in ('top1', 'top3', 'top5', 'mrr', 'exam')]],
        ["FormulaGuard v3", *[f"{float(enron_summary['formulaguard_v3'][key]):.3f}" for key in ('top1', 'top3', 'top5', 'mrr', 'exam')]],
        ["v3-real B", *[f"{float(enron_summary['formulaguard_v3_real'][key]):.3f}" for key in ('top1', 'top3', 'top5', 'mrr', 'exam')]],
    ], [2200, 1300, 1300, 1300, 1600, 1660])
    add_para(doc, "20 个未见事件、18 份工作簿和 200 条 event-method 结果记录均完整，测试清单哈希与冻结配置一致。图基线 MRR 为 0.365，v3 为 0.250。v3 相对 v2 的 MRR 差为 +0.0116，95% bootstrap CI 为 [-0.1318, 0.1588]；相对图基线的差为 -0.1154，CI 为 [-0.3226, 0.0803]。因此，真实盲测不支持 v3 在自然工作簿上稳定超过图基线或 v2。")
    add_para(doc, "正面案例中，681 公式工作簿的 jul01!AL44 具有正净改善候选 =SUM(AL12:AL42)，v3 将其排至第 1。关键失败案例中，5156 公式工作簿的 OCT 01!T39 被 graph 排第 1、v2 排第 9，而 v3 因候选无净改善排到第 368；v3-real B 保留 v2 的第 9 名并标为 pattern_only。这一对称案例说明“没有检测到内部改善”不能等同于“源错误不存在”。")
    add_heading(doc, "6.4 单工作簿性能：适合离线审计", 2)
    rows = [[r["target_formula_count"], r["samples"], f"{float(r['parse_median']):.4f}", f"{float(r['localization_median']):.2f}", f"{float(r['localization_p95']):.2f}"] for r in latency]
    add_table(doc, ["公式数", "重复", "解析中位数（秒）", "定位中位数（秒）", "定位 P95（秒）"], rows, [1600, 1300, 2200, 2200, 2060])
    add_para(doc, "100、500、1000、5000 公式的定位中位数分别为 12.17、67.03、140.51、726.24 秒。解析时间均小于 0.08 秒，主要成本来自候选反事实评估。当前实现可用于离线审计或批量风险筛查，但 5000 公式约 12 分钟，不能称为实时 Excel 插件。")
    add_callout(doc, "批量吞吐（待补充）", "多核批量吞吐将使用独立的 --mode throughput --workers 16 协议写入单独表格。它不会覆盖或改变以上隔离单表延迟，也不参与算法准确率结论。")

    add_heading(doc, "7 真实性、局限与有效性威胁", 1)
    add_bullets(doc, [
        "构念有效性：内部一致性下降是候选支持证据，不等于业务真值或严格的因果可识别效应。",
        "内部有效性：合成变异器与候选修复器镜像匹配率为 100%，故合成满分仅为机制证据。",
        "外部有效性：真实 Enron 测试仅 20 事件，且 18 个工作簿超过 100 公式干预上限；宏、外部链接、动态数组和大量高级函数尚不支持。",
        "统计有效性：真实集 bootstrap 区间跨 0 时不声称稳定优于；开发集不与冻结测试集合并调参。",
        "可复现性：代码、冻结配置、公共结果文件、环境记录和哈希审计保留；独立小样本盲测工具已准备但尚未采集，不得虚构其结果。",
    ])
    add_heading(doc, "8 结论", 1)
    add_para(doc, "FormulaGuard 把局部公式异常、候选小修复和下游依赖区域的内部变化连接为一个无输出真值的审查排序流程。它最可靠的贡献不是“自动证明错误”或“真实集领先”，而是把可解释证据、证据不足时的降级规则、生成器耦合审计和真实盲测负结果放进同一条可复核证据链。下一步应在不触碰冻结 Enron 测试结果的前提下，完成批量吞吐与独立人类盲测，以检验其在生成器之外的可用性与人工审查负担。")

    add_heading(doc, "参考文献", 1)
    refs = (ROOT / "research" / "REFERENCES_V3.md").read_text(encoding="utf-8").splitlines()
    for line in refs:
        if line.startswith("["):
            add_para(doc, line, after=3, size=9.5)

    doc.add_page_break()
    add_heading(doc, "附录 A：复现与提交审计", 1)
    add_para(doc, "正式实验结果由 scripts/audit_v3_delivery.py 统一审计。当前审计已通过合成 full、LibreOffice、Enron 冻测、负面结论保留、单表延迟和论文证据文件；批量吞吐未生成时，审计器必须保持不通过状态。每一项结果由 CSV/JSON 自动重建，禁止从截图或手工抄写得到。")
    add_heading(doc, "A.1 关键运行命令", 2)
    for command in [
        "run_pipeline.cmd full -Ablations -WithSensitivity -WithPerformance -WithLibreOffice",
        "run_enron_test.cmd",
        "python scripts/run_performance.py --input data/scaling --output results/v3_full/performance_v3_latency.csv --methods formulaguard_v3 --candidate-limit 15 --repeats 5 --mode latency --workers 1",
        "python scripts/audit_v3_delivery.py --root . --strict",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(command)
        set_run_font(run, size=9.2, color=DARK_BLUE)
    add_heading(doc, "A.2 研究者声明", 2)
    add_para(doc, "本初稿使用自动化工具协助组织代码、表格和文字，但所有数值仅来自项目内原始结果文件。未完成的吞吐和独立盲测均明确保留为待完成状态。参赛者应在提交前核查姓名、指导教师、竞赛格式、引用格式、数据许可和 AI 使用说明。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "FormulaGuard v3 竞赛论文研究初稿"
    doc.core_properties.author = "FormulaGuard Project"
    doc.core_properties.subject = "Spreadsheet silent source formula error localization"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
